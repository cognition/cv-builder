"""Flask API tests for drafts, match, variants, and variant deletion."""

from __future__ import annotations

import runpy
import sys
from pathlib import Path
from typing import TYPE_CHECKING, Any, Generator

import pytest
from ruamel.yaml import YAML

from cvbuilder.database import SnippetDatabase
from cvbuilder.document_store import DocumentStore
from cvbuilder.models import DetailLevel, Snippet, SnippetVariant

if TYPE_CHECKING:
    from _pytest.capture import CaptureFixture
    from _pytest.fixtures import FixtureRequest
    from _pytest.logging import LogCaptureFixture
    from _pytest.monkeypatch import MonkeyPatch
    from pytest_mock.plugin import MockerFixture
    from flask.testing import FlaskClient


def _master_text(api_app: dict[str, Any]) -> str:
    """Return the master YAML text from the test document store."""
    document = api_app["document_store"].get_master()
    assert document is not None
    return document.content_yaml


def _master_data(api_app: dict[str, Any]) -> Any:
    """Parse the master YAML text from the test document store."""
    return YAML().load(_master_text(api_app))


@pytest.fixture
def api_app(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, repo_fixture: Path
) -> Generator[Any, None, None]:
    """Load serve-editor with an isolated DB and variants directory.

    Args:
        tmp_path: Temporary directory for the DB.
        monkeypatch: Pytest monkeypatch fixture.
        repo_fixture: Minimal repository tree.

    Yields:
        The Flask application module namespace.
    """
    db_path = tmp_path / "api.db"
    monkeypatch.setenv("SNIPPETS_DB", str(db_path))

    # Point cvweb.REPO_ROOT / WEB_DIR / variants at the fixture tree.
    scripts = Path(__file__).resolve().parents[1] / "scripts"
    src = Path(__file__).resolve().parents[1] / "src"
    monkeypatch.syspath_prepend(str(scripts))
    monkeypatch.syspath_prepend(str(src))

    import cvweb

    monkeypatch.setattr(cvweb, "REPO_ROOT", repo_fixture)
    monkeypatch.setattr(cvweb, "WEB_DIR", repo_fixture / "cv" / "web")
    monkeypatch.setattr(cvweb, "DATA_FILE", repo_fixture / "cv" / "web" / "data.yaml")

    # Provide minimal builder/variants HTML for page routes.
    (repo_fixture / "cv" / "web" / "builder.html").write_text(
        "<html>builder</html>", encoding="utf-8"
    )
    (repo_fixture / "cv" / "web" / "variants.html").write_text(
        "<html>variants</html>", encoding="utf-8"
    )
    (repo_fixture / "cv" / "variants").mkdir(parents=True, exist_ok=True)

    # Ensure a clean import of serve-editor under the patched paths.
    sys.modules.pop("serve-editor", None)
    ns = runpy.run_path(str(scripts / "serve-editor.py"))
    monkeypatch.setattr(ns["cvweb"], "REPO_ROOT", repo_fixture)
    monkeypatch.setattr(
        ns["cvweb"], "WEB_DIR", repo_fixture / "cv" / "web"
    )
    monkeypatch.setattr(
        ns["cvweb"], "DATA_FILE", repo_fixture / "cv" / "web" / "data.yaml"
    )
    # Patch route globals so variants/assets paths resolve inside the fixture tree.
    for func_name in (
        "api_list_variants",
        "api_delete_variant_folder",
        "api_render_variant",
        "api_list_images",
        "api_upload_image",
        "api_fetch_image",
        "api_delete_image",
        "api_upload_import",
    ):
        func = ns[func_name]
        func.__globals__["VARIANTS_DIR"] = repo_fixture / "cv" / "variants"
        func.__globals__["ASSETS_DIR"] = repo_fixture / "assets" / "images"
        func.__globals__["IMPORTS_DIR"] = repo_fixture / "data" / "imports"
        func.__globals__["STAGING_DIR"] = repo_fixture / "data" / "imports" / "staging"
        func.__globals__["cvweb"] = ns["cvweb"]

    database = SnippetDatabase(db_path)
    database.ensure_schema()
    sid = database.create_snippet(
        Snippet(
            category="skill",
            heading="Python",
            tags=["python"],
        )
    )
    database.upsert_variant(
        SnippetVariant(
            snippet_id=sid,
            detail_level=DetailLevel.STANDARD.value,
            content="Python development",
        )
    )
    document_store = DocumentStore(database)
    document_store.bootstrap_from_filesystem(repo_fixture)
    ns["document_store"] = document_store
    ns["data_file"] = repo_fixture / "cv" / "web" / "data.yaml"
    ns["client"] = ns["app"].test_client()
    yield ns


@pytest.fixture
def client(api_app: dict[str, Any]) -> "FlaskClient":
    """Return a Flask test client for the loaded app."""
    app = api_app["app"]
    app.config["TESTING"] = True
    return app.test_client()


class TestApiEndpoints:
    """Exercise the new Flask endpoints with a test client."""

    def test_drafts_round_trip(self, client: "FlaskClient") -> None:
        """PUT/GET/DELETE drafts should succeed."""
        put = client.put(
            "/api/drafts/demo",
            json={"selections": [{"snippet_id": 1, "detail_level": "standard"}]},
        )
        assert put.status_code == 200
        assert put.get_json()["name"] == "demo"
        get = client.get("/api/drafts/demo")
        assert get.status_code == 200
        listed = client.get("/api/drafts")
        assert any(item["name"] == "demo" for item in listed.get_json())
        deleted = client.delete("/api/drafts/demo")
        assert deleted.status_code == 200
        missing = client.get("/api/drafts/demo")
        assert missing.status_code == 404

    def test_put_draft_with_apply_updates_working_not_variants(
        self, client: "FlaskClient"
    ) -> None:
        """PUT apply=true updates Working Draft and does not add variants."""
        import os

        store = DocumentStore(SnippetDatabase(Path(os.environ["SNIPPETS_DB"])))
        before = [item.name for item in store.list_variants()]
        snippets = client.get("/api/snippets").get_json()
        snippet_id = snippets[0]["id"]
        selections = [
            {
                "snippet_id": snippet_id,
                "detail_level": "standard",
                "section": "skill",
            }
        ]
        resp = client.put(
            "/api/drafts/demo",
            json={"selections": selections, "apply": True},
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["applied"] is True
        working = store.get_working()
        assert working is not None
        assert "Python development" in working.content_yaml
        assert [item.name for item in store.list_variants()] == before

    def test_post_draft_apply_reapplies_saved_selections(
        self, client: "FlaskClient"
    ) -> None:
        """POST /api/drafts/<name>/apply rebuilds Working Draft from draft."""
        import os

        snippets = client.get("/api/snippets").get_json()
        snippet_id = snippets[0]["id"]
        selections = [
            {
                "snippet_id": snippet_id,
                "detail_level": "standard",
                "section": "skill",
            }
        ]
        saved = client.put("/api/drafts/demo", json={"selections": selections})
        assert saved.status_code == 200
        resp = client.post("/api/drafts/demo/apply", json={})
        assert resp.status_code == 200
        assert resp.get_json()["ok"] is True
        store = DocumentStore(SnippetDatabase(Path(os.environ["SNIPPETS_DB"])))
        working = store.get_working()
        assert working is not None
        assert "Python development" in working.content_yaml

    def test_working_draft_add_snippets_merges_without_wipe(
        self, client: "FlaskClient"
    ) -> None:
        """POST /api/working-draft/add-snippets appends into the Working Draft."""
        import os

        store = DocumentStore(SnippetDatabase(Path(os.environ["SNIPPETS_DB"])))
        store.upsert_working(
            "person:\n  first_name: Test\n"
            "bio:\n  - Keep me.\n"
            "skills:\n  technical: []\n  functional: []\n"
            "experience: []\neducation: []\n"
        )
        snippets = client.get("/api/snippets").get_json()
        snippet_id = snippets[0]["id"]
        resp = client.post(
            "/api/working-draft/add-snippets",
            json={
                "selections": [
                    {
                        "snippet_id": snippet_id,
                        "detail_level": "standard",
                        "section": "skill",
                    }
                ]
            },
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["added_count"] == 1
        working = store.get_working()
        assert working is not None
        assert "Keep me." in working.content_yaml
        assert "Python development" in working.content_yaml

    def test_working_draft_conflicts_resolve_keep_both(
        self, client: "FlaskClient"
    ) -> None:
        """Conflict endpoints list highlights and clear them on keep_both."""
        import os

        from cvbuilder.models import DetailLevel, Snippet, SnippetVariant

        database = SnippetDatabase(Path(os.environ["SNIPPETS_DB"]))
        store = DocumentStore(database)
        store.upsert_working(
            "person:\n  first_name: Test\n"
            "bio: []\n"
            "skills:\n  technical: []\n  functional: []\n"
            "experience: []\neducation: []\n"
        )
        snippet_id = database.create_snippet(
            Snippet(
                category="bio",
                heading="Intro",
                tags=["api"],
                content_hash="api-conflict",
            )
        )
        database.upsert_variant(
            SnippetVariant(
                snippet_id=snippet_id,
                detail_level=DetailLevel.BRIEF.value,
                content="API brief bio.",
            )
        )
        database.upsert_variant(
            SnippetVariant(
                snippet_id=snippet_id,
                detail_level=DetailLevel.STANDARD.value,
                content="API standard bio.",
            )
        )
        first = client.post(
            "/api/working-draft/add-snippets",
            json={
                "selections": [
                    {
                        "snippet_id": snippet_id,
                        "detail_level": "standard",
                        "section": "bio",
                    }
                ]
            },
        )
        assert first.status_code == 200
        second = client.post(
            "/api/working-draft/add-snippets",
            json={
                "selections": [
                    {
                        "snippet_id": snippet_id,
                        "detail_level": "brief",
                        "section": "bio",
                    }
                ]
            },
        )
        assert second.status_code == 200
        body = second.get_json()
        assert body["warning"]
        listed = client.get("/api/working-draft/conflicts")
        assert listed.status_code == 200
        assert len(listed.get_json()["conflicts"]) >= 2
        resolved = client.post(
            "/api/working-draft/conflicts/resolve",
            json={"action": "keep_both"},
        )
        assert resolved.status_code == 200
        remaining = client.get("/api/working-draft/conflicts")
        assert remaining.get_json()["conflicts"] == []

    def test_load_variant_into_working_draft(
        self, client: "FlaskClient"
    ) -> None:
        """POST /api/working-draft/load-variant replaces Working Draft content."""
        import os

        store = DocumentStore(SnippetDatabase(Path(os.environ["SNIPPETS_DB"])))
        store.upsert_working(
            "person:\n  first_name: KeepMe\n  last_name: Please\n"
            "bio:\n  - Stale bio.\n"
            "skills:\n  technical: []\n  functional: []\n"
            "experience: []\neducation: []\n"
        )
        store.upsert_variant(
            "ready-one",
            "person:\n  first_name: Ignore\n"
            "bio:\n  - Loaded from version.\n"
            "skills:\n  technical:\n    - Loaded skill\n  functional: []\n"
            "experience: []\neducation: []\n",
        )
        resp = client.post(
            "/api/working-draft/load-variant",
            json={"name": "ready-one"},
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["ok"] is True
        assert body["name"] == "ready-one"
        working = store.get_working()
        assert working is not None
        assert "KeepMe" in working.content_yaml
        assert "Loaded from version." in working.content_yaml
        assert "Loaded skill" in working.content_yaml
        assert "Stale bio." not in working.content_yaml
        assert "Ignore" not in working.content_yaml

    def test_match_endpoint(self, client: "FlaskClient") -> None:
        """POST /api/match should return ranked hits for known terms."""
        resp = client.post("/api/match", json={"text": "Need strong Python skills"})
        assert resp.status_code == 200
        body = resp.get_json()
        assert body
        assert body[0]["snippet_id"]
        assert "python" in body[0]["matched_terms"]

    def test_compose_defaults_to_database_variant(
        self, client: "FlaskClient", api_app: dict[str, Any], repo_fixture: Path
    ) -> None:
        """POST /api/compose should save a DB variant without exporting files."""
        response = client.post(
            "/api/compose",
            json={
                "name": "api-compose",
                "selections": [
                    {
                        "snippet_id": 1,
                        "detail_level": "standard",
                        "section": "skill",
                    }
                ],
            },
        )

        assert response.status_code == 200
        body = response.get_json()
        assert body["data_yaml"] is None
        assert body["pdf"] is None
        assert (
            repo_fixture / "cv" / "variants" / "api-compose" / "data.yaml"
        ).exists() is False
        variant = api_app["document_store"].get_variant("api-compose")
        assert variant is not None
        assert "Python development" in variant.content_yaml

    def test_delete_variant_level(self, client: "FlaskClient") -> None:
        """DELETE /api/snippets/<id>/variants/<level> should remove one level."""
        created = client.post(
            "/api/snippets",
            json={
                "category": "part",
                "heading": "Temp",
                "variants": {
                    "brief": "short",
                    "standard": "normal",
                },
            },
        )
        assert created.status_code == 201
        snippet_id = created.get_json()["id"]
        deleted = client.delete(f"/api/snippets/{snippet_id}/variants/brief")
        assert deleted.status_code == 200
        fetched = client.get(f"/api/snippets/{snippet_id}")
        levels = {v["detail_level"] for v in fetched.get_json()["variants"]}
        assert "brief" not in levels
        assert "standard" in levels

    def test_save_updates_database_not_file(self, api_app: dict[str, Any]) -> None:
        """Saving an edit must mutate cv_documents, not the fixture YAML path."""
        client = api_app["client"]
        yaml_path = api_app["data_file"]
        before_file = yaml_path.read_text(encoding="utf-8")
        master_before = _master_text(api_app)

        response = client.post(
            "/api/save",
            json=[{"path": "bio[0]", "value": "Saved only in DB"}],
        )

        assert response.status_code == 200
        assert yaml_path.read_text(encoding="utf-8") == before_file
        assert "Saved only in DB" in _master_text(api_app)
        assert _master_text(api_app) != master_before

    def test_pin_create_list_restore_and_delete(self, api_app: dict[str, Any]) -> None:
        """Pin APIs should snapshot, restore, list, and delete master versions."""
        client = api_app["client"]
        store = api_app["document_store"]
        before = _master_text(api_app)

        created = client.post("/api/pins", json={"document": "master", "label": "v1"})

        assert created.status_code == 200
        pin_id = created.get_json()["id"]
        listed = client.get("/api/pins?document=master")
        assert listed.status_code == 200
        assert [pin["label"] for pin in listed.get_json()] == ["v1"]

        saved = client.post(
            "/api/save", json=[{"path": "bio[0]", "value": "After pin"}]
        )
        assert saved.status_code == 200
        changed = store.get_master()
        assert changed is not None
        assert "After pin" in changed.content_yaml

        restored = client.post(f"/api/pins/{pin_id}/restore")

        assert restored.status_code == 200
        master = store.get_master()
        assert master is not None
        assert master.content_yaml == before
        assert master.id is not None
        labels = [pin.label for pin in store.list_pins(master.id)]
        assert f"before-restore:{pin_id}" in labels

        deleted = client.delete(f"/api/pins/{pin_id}")
        assert deleted.status_code == 200
        listed_after_delete = client.get("/api/pins?document=master").get_json()
        remaining = [pin["id"] for pin in listed_after_delete]
        assert pin_id not in remaining

    def test_structure_replace_text(
        self, client: "FlaskClient", api_app: dict[str, Any]
    ) -> None:
        """POST /api/structure op=replace should overwrite a leaf value."""
        before_file = api_app["data_file"].read_text(encoding="utf-8")
        resp = client.post(
            "/api/structure",
            json={"op": "replace", "path": "bio[0]", "value": "Replacement bio."},
        )
        assert resp.status_code == 200
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file
        data = _master_data(api_app)
        assert str(data["bio"][0]).strip() == "Replacement bio."

    def test_structure_replace_subsection(
        self, client: "FlaskClient", api_app: dict[str, Any]
    ) -> None:
        """op=replace-subsection should rebuild heading/paragraphs/bullets."""
        before_file = api_app["data_file"].read_text(encoding="utf-8")
        resp = client.post(
            "/api/structure",
            json={
                "op": "replace-subsection",
                "path": "experience[0].subsections[0]",
                "heading": "Swapped Heading",
                "content": "Intro paragraph text.\n\n- First bullet\n- Second bullet",
            },
        )
        assert resp.status_code == 200
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file
        data = _master_data(api_app)
        sub = data["experience"][0]["subsections"][0]
        assert sub["heading"] == "Swapped Heading"
        assert list(sub["bullets"]) == ["First bullet", "Second bullet"]
        assert list(sub["paragraphs"]) == ["Intro paragraph text."]

    def test_structure_replace_requires_value(
        self, client: "FlaskClient"
    ) -> None:
        """op=replace without a usable value should fail with 400."""
        resp = client.post(
            "/api/structure",
            json={"op": "replace", "path": "bio[0]", "value": "  "},
        )
        assert resp.status_code == 400

    def test_structure_then_undo_redo(
        self, client: "FlaskClient", api_app: dict[str, Any]
    ) -> None:
        """Structural changes should be reversible via /api/undo and /api/redo."""
        before = _master_text(api_app)
        before_file = api_app["data_file"].read_text(encoding="utf-8")
        inserted = client.post(
            "/api/structure",
            json={"op": "insert", "path": "bio", "value": "Undo me"},
        )
        assert inserted.status_code == 200
        assert inserted.get_json()["can_undo"] is True
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file

        status = client.get("/api/history")
        assert status.status_code == 200
        assert status.get_json()["can_undo"] is True

        undone = client.post("/api/undo")
        assert undone.status_code == 200
        assert _master_text(api_app) == before
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file
        assert undone.get_json()["can_redo"] is True

        redone = client.post("/api/redo")
        assert redone.status_code == 200
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file
        data = _master_data(api_app)
        assert "Undo me" in [str(part).strip() for part in data["bio"]]

    def test_empty_save_skips_history(self, client: "FlaskClient") -> None:
        """An empty save payload should not create an undo entry."""
        resp = client.post("/api/save", json=[])
        assert resp.status_code == 200
        assert resp.get_json()["can_undo"] is False

    def test_export_yaml_writes_file_only_when_requested(
        self, api_app: dict[str, Any], tmp_path: Path
    ) -> None:
        """YAML export should write DB-backed content to the requested path."""
        out = tmp_path / "exported.yaml"
        before_file = api_app["data_file"].read_text(encoding="utf-8")
        api_app["document_store"].upsert_master(
            "person:\n  first_name: DB\n  last_name: Export\nbio:\n  - From DB\n"
        )

        response = api_app["client"].post(
            "/api/export",
            json={"format": "yaml", "path": str(out)},
        )

        assert response.status_code == 200
        assert out.is_file()
        assert "From DB" in out.read_text(encoding="utf-8")
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file

    def test_export_yaml_write_failure_removes_partial_target(
        self,
        api_app: dict[str, Any],
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        """A failed export should not leave a corrupt partial target behind."""
        out = tmp_path / "partial.yaml"
        original_write_text = Path.write_text

        def failing_write_text(
            path: Path,
            data: str,
            encoding: str | None = None,
            errors: str | None = None,
            newline: str | None = None,
        ) -> int:
            """Write a partial file at the target path, then fail."""
            if path == out:
                path.parent.mkdir(parents=True, exist_ok=True)
                path.write_bytes(b"partial export")
                raise OSError("simulated write failure")
            return original_write_text(
                path, data, encoding=encoding, errors=errors, newline=newline
            )

        monkeypatch.setattr(Path, "write_text", failing_write_text)

        response = api_app["client"].post(
            "/api/export",
            json={"format": "yaml", "path": str(out)},
        )

        assert response.status_code == 500
        assert out.exists() is False

    def test_export_markdown_writes_variant_by_name(
        self, api_app: dict[str, Any], tmp_path: Path
    ) -> None:
        """Markdown export should support named DB-backed variant documents."""
        out = tmp_path / "variant.md"
        api_app["document_store"].upsert_variant(
            "plant-role",
            "person:\n  first_name: Homer\n  last_name: Simpson\nbio:\n  - Variant bio\n",
        )

        response = api_app["client"].post(
            "/api/export",
            json={
                "format": "markdown",
                "document": "variant",
                "name": "plant-role",
                "path": str(out),
            },
        )

        assert response.status_code == 200
        assert "# Homer Simpson" in out.read_text(encoding="utf-8")
        assert "Variant bio" in out.read_text(encoding="utf-8")

    def test_export_default_pdf_uses_database_without_yaml_side_effect(
        self,
        api_app: dict[str, Any],
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        """Default export should render the DB master as PDF without YAML writes."""
        captured: dict[str, Any] = {}
        before_file = api_app["data_file"].read_text(encoding="utf-8")
        api_app["document_store"].upsert_master(
            "person:\n  first_name: PDF\n  last_name: Export\nbio:\n  - PDF from DB\n"
        )

        def fake_export_pdf(out_pdf: Path, data: dict[str, Any]) -> None:
            """Capture PDF export inputs without launching Chrome."""
            captured["out_pdf"] = out_pdf
            captured["data"] = data
            out_pdf.parent.mkdir(parents=True, exist_ok=True)
            out_pdf.write_bytes(b"%PDF-1.4 fake")

        monkeypatch.setitem(
            api_app["api_export"].__globals__,
            "PREVIEW_PDF",
            tmp_path / "cv.pdf",
        )
        monkeypatch.setattr(api_app["cvweb"], "export_pdf", fake_export_pdf)

        response = api_app["client"].post("/api/export", json={})

        assert response.status_code == 200
        assert captured["data"]["bio"] == ["PDF from DB"]
        assert captured["out_pdf"] == tmp_path / "cv.pdf"
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file

    def test_edit_page_without_master_returns_404(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """GET /edit should 404 when bootstrap finds no Master CV."""
        repo_root = tmp_path / "repo"
        web = repo_root / "cv" / "web"
        web.mkdir(parents=True)
        (repo_root / "cv" / "variants").mkdir(parents=True)

        db_path = tmp_path / "api.db"
        monkeypatch.setenv("SNIPPETS_DB", str(db_path))

        scripts = Path(__file__).resolve().parents[1] / "scripts"
        src = Path(__file__).resolve().parents[1] / "src"
        real_web = Path(__file__).resolve().parents[1] / "cv" / "web"
        monkeypatch.syspath_prepend(str(scripts))
        monkeypatch.syspath_prepend(str(src))

        import cvweb

        monkeypatch.setattr(cvweb, "REPO_ROOT", repo_root)
        monkeypatch.setattr(cvweb, "WEB_DIR", real_web)
        monkeypatch.setattr(cvweb, "DATA_FILE", web / "data.yaml")

        sys.modules.pop("serve-editor", None)
        ns = runpy.run_path(str(scripts / "serve-editor.py"))
        monkeypatch.setattr(ns["cvweb"], "REPO_ROOT", repo_root)
        monkeypatch.setattr(ns["cvweb"], "WEB_DIR", real_web)
        monkeypatch.setattr(
            ns["cvweb"], "DATA_FILE", web / "data.yaml"
        )

        database = SnippetDatabase(db_path)
        database.ensure_schema()
        document_store = DocumentStore(database)
        document_store.bootstrap_from_filesystem(repo_root)
        assert document_store.get_master() is None

        app = ns["app"]
        app.config["TESTING"] = True
        client = app.test_client()

        edit = client.get("/edit")
        assert edit.status_code == 404
        assert b"Working Draft unavailable" in edit.data

        person = client.get("/api/person")
        assert person.status_code == 404
        assert "Master CV document is not available" in person.get_json()["error"]

    def test_studio_routes_are_top_level(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch, repo_fixture: Path
    ) -> None:
        """Studio pages should live at the top path, not under /cv/web."""
        db_path = tmp_path / "api.db"
        monkeypatch.setenv("SNIPPETS_DB", str(db_path))

        scripts = Path(__file__).resolve().parents[1] / "scripts"
        src = Path(__file__).resolve().parents[1] / "src"
        real_web = Path(__file__).resolve().parents[1] / "cv" / "web"
        monkeypatch.syspath_prepend(str(scripts))
        monkeypatch.syspath_prepend(str(src))

        import cvweb

        monkeypatch.setattr(cvweb, "REPO_ROOT", repo_fixture)
        monkeypatch.setattr(cvweb, "WEB_DIR", real_web)
        monkeypatch.setattr(
            cvweb, "DATA_FILE", repo_fixture / "cv" / "web" / "data.yaml"
        )

        sys.modules.pop("serve-editor", None)
        ns = runpy.run_path(str(scripts / "serve-editor.py"))
        monkeypatch.setattr(ns["cvweb"], "REPO_ROOT", repo_fixture)
        monkeypatch.setattr(ns["cvweb"], "WEB_DIR", real_web)
        monkeypatch.setattr(
            ns["cvweb"], "DATA_FILE", repo_fixture / "cv" / "web" / "data.yaml"
        )

        database = SnippetDatabase(db_path)
        database.ensure_schema()
        DocumentStore(database).bootstrap_from_filesystem(repo_fixture)

        app = ns["app"]
        app.config["TESTING"] = True
        client = app.test_client()

        home = client.get("/")
        assert home.status_code == 200
        assert b'href="/edit"' in home.data
        assert b'href="/cv/web/edit"' not in home.data

        edit = client.get("/edit")
        assert edit.status_code == 200
        assert b"cv-document" in edit.data

        theme = client.get("/src/theme.css")
        assert theme.status_code == 200
        assert b":root" in theme.data or b"--" in theme.data

        legacy = client.get("/cv/web/edit", follow_redirects=False)
        assert legacy.status_code == 301
        assert legacy.headers["Location"].endswith("/edit")

    def test_home_lists_variants_when_exports_live_outside_repo(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch, repo_fixture: Path
    ) -> None:
        """Home must not 500 when variant PDFs live under CV_DATA_ROOT."""
        import os

        data_root = tmp_path / "data-root"
        variants_dir = data_root / "cv" / "variants"
        sample_dir = variants_dir / "sample"
        sample_dir.mkdir(parents=True)
        (sample_dir / "sample.pdf").write_bytes(b"%PDF-1.4")
        monkeypatch.setenv("CV_DATA_ROOT", str(data_root))
        monkeypatch.setenv("SNIPPETS_DB", str(data_root / "snippets.db"))
        monkeypatch.setenv("VARIANTS_DIR", str(variants_dir))

        scripts = Path(__file__).resolve().parents[1] / "scripts"
        src = Path(__file__).resolve().parents[1] / "src"
        monkeypatch.syspath_prepend(str(scripts))
        monkeypatch.syspath_prepend(str(src))
        import cvweb

        monkeypatch.setattr(cvweb, "REPO_ROOT", repo_fixture)
        monkeypatch.setattr(cvweb, "WEB_DIR", repo_fixture / "cv" / "web")
        monkeypatch.setattr(
            cvweb, "DATA_FILE", repo_fixture / "cv" / "web" / "data.yaml"
        )
        sys.modules.pop("serve-editor", None)
        ns = runpy.run_path(str(scripts / "serve-editor.py"))
        monkeypatch.setattr(ns["cvweb"], "REPO_ROOT", repo_fixture)
        monkeypatch.setattr(ns["cvweb"], "WEB_DIR", repo_fixture / "cv" / "web")
        for name in (
            "_list_variants",
            "_variant_file_url",
            "home_page",
            "api_list_variants",
            "serve_variant_export",
            "api_render_variant",
        ):
            ns[name].__globals__["VARIANTS_DIR"] = variants_dir
        ns["VARIANTS_DIR"] = variants_dir

        database = SnippetDatabase(Path(os.environ["SNIPPETS_DB"]))
        database.ensure_schema()
        store = DocumentStore(database)
        store.bootstrap_from_filesystem(repo_fixture)
        store.upsert_variant("sample", "person:\n  first_name: Sample\n")
        # Make sure home has a Working Draft document available.
        if store.get_working() is None:
            store.upsert_working(
                "person:\n  first_name: Test\n  last_name: User\n"
                "bio: []\nskills:\n  technical: []\n  functional: []\n"
                "experience: []\neducation: []\n"
            )
        ns["_document_store"] = lambda: store  # type: ignore[assignment]
        ns["home_page"].__globals__["_document_store"] = lambda: store
        ns["api_list_variants"].__globals__["_document_store"] = lambda: store
        ns["_list_variants"].__globals__["_document_store"] = lambda: store
        ns["app"].config["TESTING"] = True
        client = ns["app"].test_client()

        listed = client.get("/api/variants")
        assert listed.status_code == 200
        body = listed.get_json()
        assert any(item["name"] == "sample" for item in body)
        sample = next(item for item in body if item["name"] == "sample")
        assert sample["pdf"] == "cv/variants/sample/sample.pdf"

        # Direct call mirrors what home_page uses — must not raise ValueError.
        variants = ns["_list_variants"]()
        assert any(
            item["name"] == "sample" and item["pdf"] == "cv/variants/sample/sample.pdf"
            for item in variants
        )

        pdf = client.get("/cv/variants/sample/sample.pdf")
        assert pdf.status_code == 200
        assert pdf.data.startswith(b"%PDF")

    def test_variants_list_and_delete(
        self, client: "FlaskClient", api_app: dict[str, Any], repo_fixture: Path
    ) -> None:
        """Variant list/delete should operate on DB-backed variant documents."""
        variant_dir = repo_fixture / "cv" / "variants" / "sample"
        variant_dir.mkdir(parents=True, exist_ok=True)
        (variant_dir / "data.yaml").write_text("person: {}\n", encoding="utf-8")
        (variant_dir / "sample.pdf").write_bytes(b"%PDF-1.4")
        api_app["document_store"].upsert_variant(
            "sample",
            "person:\n  first_name: Variant\n  last_name: Person\n",
        )

        listed = client.get("/api/variants")
        assert listed.status_code == 200
        names = {item["name"] for item in listed.get_json()}
        assert "sample" in names

        deleted = client.delete("/api/variants/sample")
        assert deleted.status_code == 200
        assert api_app["document_store"].get_variant("sample") is None
        assert not variant_dir.exists()

    def test_render_variant_uses_database_document(
        self,
        client: "FlaskClient",
        api_app: dict[str, Any],
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        """Variant rendering should export a PDF from the DB document content."""
        api_app["document_store"].upsert_variant(
            "db-render",
            "\n".join(
                [
                    "person:",
                    "  first_name: Render",
                    "  last_name: Person",
                    "bio:",
                    "  - Rendered bio",
                    "",
                ]
            ),
        )

        def _fake_export(pdf_path: Path, data: Any = None) -> None:
            """Skip Chrome and write a tiny PDF marker."""
            assert isinstance(data, dict)
            assert data["bio"] == ["Rendered bio"]
            pdf_path.parent.mkdir(parents=True, exist_ok=True)
            pdf_path.write_bytes(b"%PDF-1.4 fake")

        monkeypatch.setattr(api_app["cvweb"], "export_pdf", _fake_export)

        rendered = client.post("/api/variants/db-render/render")

        assert rendered.status_code == 200
        body = rendered.get_json()
        assert body["ok"] is True
        assert body["pdf"] == "cv/variants/db-render/db-render.pdf"

    def test_image_upload_and_list(
        self, client: "FlaskClient", repo_fixture: Path
    ) -> None:
        """Uploading an image should store it and list it afterwards."""
        from io import BytesIO

        png_bytes = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
        )
        upload = client.post(
            "/api/images/upload",
            data={"file": (BytesIO(png_bytes), "test icon.png")},
            content_type="multipart/form-data",
        )
        assert upload.status_code == 201
        body = upload.get_json()
        assert body["name"].endswith(".png")
        assert body["data_path"].startswith("../../assets/images/")
        stored = repo_fixture / "assets" / "images" / body["name"]
        assert stored.is_file()

        listed = client.get("/api/images")
        assert listed.status_code == 200
        names = {item["name"] for item in listed.get_json()}
        assert body["name"] in names

    def test_image_delete_removes_file_and_clears_photo(
        self, client: "FlaskClient", repo_fixture: Path
    ) -> None:
        """Deleting an uploaded image removes the file and clears person.photo."""
        from io import BytesIO

        png_bytes = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"
            b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
        )
        upload = client.post(
            "/api/images/upload",
            data={"file": (BytesIO(png_bytes), "nuke.png")},
            content_type="multipart/form-data",
        )
        assert upload.status_code == 201
        name = upload.get_json()["name"]
        data_path = upload.get_json()["data_path"]
        saved = client.post(
            "/api/save",
            json=[{"path": "person.photo", "value": data_path}],
        )
        assert saved.status_code == 200

        deleted = client.delete(f"/api/images/{name}")
        assert deleted.status_code == 200
        assert deleted.get_json()["ok"] is True
        assert deleted.get_json()["cleared_profile_photo"] is True
        assert not (repo_fixture / "assets" / "images" / name).exists()
        listed = {item["name"] for item in client.get("/api/images").get_json()}
        assert name not in listed
        person = client.get("/api/person").get_json()
        assert person.get("photo") in ("", None)

    def test_image_delete_rejects_path_traversal(self, client: "FlaskClient") -> None:
        """Path traversal must not delete files outside the assets directory."""
        resp = client.delete("/api/images/../../etc/passwd")
        assert resp.status_code in {400, 404}

    def test_image_upload_rejects_bad_type(self, client: "FlaskClient") -> None:
        """Non-image extensions should be rejected."""
        from io import BytesIO

        upload = client.post(
            "/api/images/upload",
            data={"file": (BytesIO(b"hello"), "notes.txt")},
            content_type="multipart/form-data",
        )
        assert upload.status_code == 400

    def test_image_fetch_rejects_bad_scheme(self, client: "FlaskClient") -> None:
        """Only http(s) URLs should be accepted for image grabs."""
        resp = client.post(
            "/api/images/fetch", json={"url": "file:///etc/passwd"}
        )
        assert resp.status_code == 400
        missing = client.post("/api/images/fetch", json={})
        assert missing.status_code == 400


class TestQuestionApiEndpoints:
    """Exercise the application-questions endpoints end-to-end."""

    def test_create_source_extracts_questions(self, client: "FlaskClient") -> None:
        """Creating a source with pasted text should extract questions."""
        created = client.post(
            "/api/question-sources",
            json={
                "title": "Government screening",
                "source_type": "form",
                "text": "Do you have a valid work permit?\nAre you willing to relocate?",
            },
        )
        assert created.status_code == 201
        body = created.get_json()
        assert body["question_count"] == 2
        source_id = body["id"]

        listed_sources = client.get("/api/question-sources")
        assert any(s["id"] == source_id for s in listed_sources.get_json())

        listed_questions = client.get(f"/api/questions?source_id={source_id}")
        prompts = {q["prompt"] for q in listed_questions.get_json()}
        assert prompts == {
            "Do you have a valid work permit?",
            "Are you willing to relocate?",
        }
        assert all(q["status"] == "needs_evidence" for q in listed_questions.get_json())

    def test_create_source_rejects_bad_type(self, client: "FlaskClient") -> None:
        """An unknown source_type should be rejected."""
        resp = client.post(
            "/api/question-sources", json={"title": "x", "source_type": "essay"}
        )
        assert resp.status_code == 400

    def test_answer_evidence_and_suggest_round_trip(
        self, client: "FlaskClient"
    ) -> None:
        """Save an answer, link evidence, then re-suggest from that evidence."""
        source = client.post(
            "/api/question-sources", json={"title": "Leadership", "source_type": "matrix"}
        ).get_json()
        question = client.post(
            "/api/question-sources", json={"title": "x2", "source_type": "matrix", "text": "Python fluency"}
        ).get_json()
        question_id = client.get(f"/api/questions?source_id={question['id']}").get_json()[0]["id"]

        # Fixture seeds a "Python" skill snippet — find its id.
        snippet_id = next(
            s["id"] for s in client.get("/api/snippets").get_json() if s["heading"] == "Python"
        )

        linked = client.post(
            f"/api/questions/{question_id}/evidence", json={"snippet_id": snippet_id}
        )
        assert linked.status_code == 201
        assert linked.get_json()["evidence"][0]["snippet_id"] == snippet_id

        after_link = client.get(f"/api/questions/{question_id}").get_json()
        assert after_link["status"] == "in_progress"

        suggested = client.post(f"/api/questions/{question_id}/suggest")
        assert suggested.status_code == 200
        assert "Python development" in suggested.get_json()["answer"]

        saved = client.put(
            f"/api/questions/{question_id}", json={"answer": "I know Python well."}
        )
        assert saved.status_code == 200
        assert saved.get_json()["status"] == "complete"

        unlinked = client.delete(f"/api/questions/{question_id}/evidence/{snippet_id}")
        assert unlinked.status_code == 200

        deleted = client.delete(f"/api/questions/{question_id}")
        assert deleted.status_code == 200
        assert client.get(f"/api/questions/{question_id}").status_code == 404

        assert client.delete(f"/api/question-sources/{source['id']}").status_code == 200

    def test_suggest_with_no_evidence_ranks_against_prompt(
        self, client: "FlaskClient"
    ) -> None:
        """With no linked evidence yet, /suggest should match against the prompt."""
        source = client.post(
            "/api/question-sources",
            json={"title": "Skills check", "source_type": "form", "text": "Python experience?"},
        ).get_json()
        question_id = client.get(f"/api/questions?source_id={source['id']}").get_json()[0]["id"]
        suggested = client.post(f"/api/questions/{question_id}/suggest")
        assert suggested.status_code == 200
        assert suggested.get_json()["evidence"]
        assert "Python development" in suggested.get_json()["answer"]


SAMPLE_RESUME_TEXT = """Homer Simpson

Summary
Platform leader with a decade of experience building resilient cloud systems.

Experience
Staff Platform Engineer — Acme Corp (2021 - Present)
- Led migration of 40+ services to Kubernetes
- Built the on-call incident response program

Skills
Kubernetes, Terraform, AWS

Education
BSc Computer Science, State University, 2013
"""


class TestResumeImportApiEndpoints:
    """Exercise the resume-import upload/review/confirm endpoints."""

    def test_upload_and_confirm_round_trip(self, client: "FlaskClient") -> None:
        from io import BytesIO

        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        assert uploaded.status_code == 201
        body = uploaded.get_json()
        assert body["file_type"] == "txt"
        assert body["counts"]["experience"] == 1
        assert body["counts"]["skills"] == 3
        assert len(body["candidates"]) == sum(body["counts"].values())
        token = body["token"]

        confirmed = client.post(f"/api/imports/{token}/confirm", json={})
        assert confirmed.status_code == 200
        confirmed_body = confirmed.get_json()
        assert confirmed_body["snippet_count"] == len(body["candidates"])

        listed = client.get("/api/imports")
        assert listed.status_code == 200
        records = listed.get_json()
        assert len(records) == 1
        assert records[0]["filename"] == "resume.txt"
        assert records[0]["snippet_count"] == confirmed_body["snippet_count"]
        import_id = records[0]["id"]

        library = client.get("/api/snippets?tag=import")
        assert len(library.get_json()) == confirmed_body["snippet_count"]

        downloaded = client.get(f"/api/imports/{import_id}/source")
        assert downloaded.status_code == 200
        assert downloaded.data == SAMPLE_RESUME_TEXT.encode()

        deleted = client.delete(f"/api/imports/{import_id}")
        assert deleted.status_code == 200
        assert client.get("/api/imports").get_json() == []

    def test_upload_rejects_unsupported_extension(self, client: "FlaskClient") -> None:
        from io import BytesIO

        resp = client.post(
            "/api/imports",
            data={"file": (BytesIO(b"whatever"), "resume.exe")},
            content_type="multipart/form-data",
        )
        assert resp.status_code == 400

    def test_confirm_respects_disabled_sections(self, client: "FlaskClient") -> None:
        from io import BytesIO

        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        token = uploaded.get_json()["token"]

        confirmed = client.post(
            f"/api/imports/{token}/confirm",
            json={
                "sections": {
                    "profile": False,
                    "experience": False,
                    "skills": True,
                    "education": False,
                }
            },
        )
        assert confirmed.status_code == 200
        assert confirmed.get_json()["snippet_count"] == 3

        skills = client.get("/api/snippets?category=skill&tag=import")
        assert len(skills.get_json()) == 3

    def test_confirm_master_mode_updates_db_and_pins(
        self, client: "FlaskClient", api_app: dict[str, Any]
    ) -> None:
        """Master imports should update DB content and create a safety pin."""
        from io import BytesIO

        before_file = api_app["data_file"].read_text(encoding="utf-8")
        before = _master_data(api_app)
        original_first = before["person"]["first_name"]

        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        token = uploaded.get_json()["token"]
        confirmed = client.post(
            f"/api/imports/{token}/confirm",
            json={"mode": "master"},
        )
        assert confirmed.status_code == 200
        body = confirmed.get_json()
        assert body["mode"] == "master"
        assert body["master_updated"] is True
        assert body["snippet_count"] > 0
        assert "backup_path" not in body
        assert api_app["data_file"].read_text(encoding="utf-8") == before_file

        after = _master_data(api_app)
        assert after["person"]["first_name"] == original_first
        assert after["bio"]  # non-empty from SAMPLE_RESUME_TEXT summary
        assert after["experience"]
        assert after["experience"][0]["company"] or after["experience"][0]["role"]
        master = api_app["document_store"].get_master()
        assert master is not None
        assert any(
            pin.label == f"before-import:{token}"
            for pin in api_app["document_store"].list_pins(master.id)
        )
        skills = client.get("/api/snippets?tag=import")
        assert len(skills.get_json()) == body["snippet_count"]

    def test_confirm_invalid_mode_returns_400(self, client: "FlaskClient") -> None:
        from io import BytesIO

        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        token = uploaded.get_json()["token"]
        resp = client.post(
            f"/api/imports/{token}/confirm",
            json={"mode": "compare"},
        )
        assert resp.status_code == 400

    def test_confirm_library_mode_default_does_not_touch_master(
        self, client: "FlaskClient"
    ) -> None:
        from io import BytesIO

        import cvweb

        before_text = cvweb.read_data_text()
        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        token = uploaded.get_json()["token"]
        confirmed = client.post(
            f"/api/imports/{token}/confirm",
            json={},
        )
        assert confirmed.status_code == 200
        body = confirmed.get_json()
        assert body["mode"] == "library"
        assert body["master_updated"] is False
        assert cvweb.read_data_text() == before_text

    def test_duplicate_candidates_are_flagged(self, client: "FlaskClient") -> None:
        from io import BytesIO

        from cvbuilder.resume_extractor import content_hash

        client.post(
            "/api/snippets",
            json={
                "category": "skill",
                "heading": "Kubernetes",
                "content_hash": content_hash("Kubernetes"),
                "content": "Kubernetes",
                "detail_level": "standard",
            },
        )

        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        candidates = uploaded.get_json()["candidates"]
        kubernetes = next(c for c in candidates if c["heading"] == "Kubernetes")
        terraform = next(c for c in candidates if c["heading"] == "Terraform")
        assert kubernetes["duplicate"] is True
        assert terraform["duplicate"] is False

    def test_upload_real_docx_through_multipart_endpoint(
        self, client: "FlaskClient"
    ) -> None:
        """A real (non-text) binary file should round-trip through the upload route."""
        from io import BytesIO

        from docx import Document

        document = Document()
        document.add_paragraph("Homer Simpson")
        document.add_paragraph("")
        document.add_paragraph("Skills")
        document.add_paragraph("Kubernetes, Terraform")
        buf = BytesIO()
        document.save(buf)
        buf.seek(0)

        uploaded = client.post(
            "/api/imports",
            data={"file": (buf, "resume.docx")},
            content_type="multipart/form-data",
        )
        assert uploaded.status_code == 201
        body = uploaded.get_json()
        assert body["file_type"] == "docx"
        assert body["counts"]["skills"] == 2

        confirmed = client.post(f"/api/imports/{body['token']}/confirm", json={})
        assert confirmed.status_code == 200
        assert confirmed.get_json()["snippet_count"] == body["counts"]["skills"] + sum(
            v for k, v in body["counts"].items() if k != "skills"
        )

        record = client.get("/api/imports").get_json()[0]
        assert record["file_type"] == "docx"
        downloaded = client.get(f"/api/imports/{record['id']}/source")
        assert downloaded.status_code == 200
        assert downloaded.data[:2] == b"PK"  # docx is a zip container

    def test_confirm_unknown_token_returns_404(self, client: "FlaskClient") -> None:
        resp = client.post("/api/imports/0011223344556677/confirm", json={})
        assert resp.status_code == 404

    def test_discard_staged_import(self, client: "FlaskClient") -> None:
        from io import BytesIO

        uploaded = client.post(
            "/api/imports",
            data={"file": (BytesIO(SAMPLE_RESUME_TEXT.encode()), "resume.txt")},
            content_type="multipart/form-data",
        )
        token = uploaded.get_json()["token"]

        discarded = client.delete(f"/api/imports/staging/{token}")
        assert discarded.status_code == 200

        confirmed = client.post(f"/api/imports/{token}/confirm", json={})
        assert confirmed.status_code == 404
